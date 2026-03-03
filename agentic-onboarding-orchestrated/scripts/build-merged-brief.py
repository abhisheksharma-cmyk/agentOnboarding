from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


PDF_SOURCE = Path(
    r"c:\Users\v00770\Downloads\VERINITE AI CUSTOMER ONBOARD APP_ System Overview and Purpose.pdf"
)


def latest_visual_docx(repo: Path) -> Path:
    matches = sorted(repo.glob("AGENTIC_ONBOARDING_VISUAL_BRIEF*.docx"), key=lambda p: p.stat().st_mtime)
    if not matches:
        raise FileNotFoundError("No AGENTIC_ONBOARDING_VISUAL_BRIEF*.docx found in repository root.")
    return matches[-1]


def read_docx_paragraphs(docx_path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    paragraphs: list[str] = []
    for p in root.findall(".//w:p", ns):
        text = "".join((node.text or "") for node in p.findall(".//w:t", ns)).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def section_between(lines: list[str], start_prefix: str, end_prefixes: tuple[str, ...]) -> list[str]:
    start_idx = next((i for i, s in enumerate(lines) if s.startswith(start_prefix)), -1)
    if start_idx == -1:
        return []
    out: list[str] = []
    for s in lines[start_idx + 1 :]:
        if any(s.startswith(prefix) for prefix in end_prefixes):
            break
        out.append(s)
    return out


def normalize_pdf_text(text: str) -> str:
    text = text.replace("\r", " ").replace("\n", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_pdf_pages(pdf_path: Path) -> list[str]:
    reader = PdfReader(str(pdf_path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(normalize_pdf_text(page.extract_text() or ""))
    return pages


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Segoe UI"
    run.font.size = Pt(20 if level == 1 else 15)
    run.font.color.rgb = RGBColor(15, 23, 42)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(31, 41, 55)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_image_if_exists(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(71, 85, 105)


def add_component_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Primary Purpose"
    hdr[2].text = "Endpoint / Port"

    rows = [
        ("Frontend (Next.js + React + TypeScript)", "Customer-facing onboarding UI, uploads, and guided chat.", "3000"),
        ("Orchestrator (Node.js + Express + TypeScript)", "Central workflow orchestration, state machine, event bus, audit trail.", "4000"),
        ("Address Verification Agent", "Address validation, normalization, confidence scoring.", "5000"),
        ("KYC Agent", "Identity and document verification workflow.", "5005"),
        ("AML Agent", "Sanctions/PEP/watchlist screening and risk flagging.", "5006"),
        ("Credit Agent", "Credit profile evaluation and eligibility signals.", "5007"),
        ("Risk + Decision Gateway", "Final policy-governed adjudication (APPROVE / DENY / ESCALATE).", "Local + orchestrator"),
    ]

    for comp, purpose, port in rows:
        cells = table.add_row().cells
        cells[0].text = comp
        cells[1].text = purpose
        cells[2].text = port


def main() -> None:
    repo = Path.cwd()
    source_docx = latest_visual_docx(repo)
    if not PDF_SOURCE.exists():
        raise FileNotFoundError(f"PDF source not found: {PDF_SOURCE}")

    lines = read_docx_paragraphs(source_docx)
    pdf_pages = extract_pdf_pages(PDF_SOURCE)

    strategic = section_between(lines, "1) Strategic Narrative", ("2) Visual Diagram A: Platform Topology",))
    yaml_deep_dive = section_between(
        lines,
        "5) Why This Is Pluggable and Configurable (YAML Deep Dive)",
        ("6) Feature Highlights with Enterprise Positioning",),
    )
    features = section_between(
        lines,
        "6) Feature Highlights with Enterprise Positioning",
        ("7) Implementation Anchors in Current Codebase",),
    )
    anchors = section_between(lines, "7) Implementation Anchors in Current Codebase", tuple())

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Segoe UI"
    normal_style.font.size = Pt(11)

    title = doc.add_paragraph()
    title_run = title.add_run("Agentic Onboarding - Consolidated Architecture and Purpose Brief")
    title_run.bold = True
    title_run.font.name = "Segoe UI"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph(
        f"Merged from: {source_docx.name} + {PDF_SOURCE.name} | Generated: {datetime.now():%Y-%m-%d %H:%M}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in meta.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(71, 85, 105)

    heading(doc, "1) System Purpose and Strategic Intent", level=1)
    body(
        doc,
        "The VERINITE AI CUSTOMER ONBOARD APP is designed as an end-to-end, AI-powered onboarding ecosystem that reduces manual processing, "
        "improves compliance execution (KYC/AML), and accelerates risk-aware customer decisioning.",
    )
    if strategic:
        for line in strategic[:2]:
            body(doc, line)
    else:
        body(
            doc,
            "The architecture is intentionally composable: specialist agents execute independently while a central policy layer preserves "
            "determinism, explainability, and operational governance.",
        )

    heading(doc, "2) Consolidated Architecture View", level=1)
    body(
        doc,
        "The combined architecture follows a microservices-plus-event-driven model: frontend requests are routed to the orchestrator, "
        "processed through slotized specialist agents, and finalized by a policy-governed Decision Gateway with traceable audit telemetry.",
    )
    add_image_if_exists(doc, repo / "documents/diagrams/architecture-topology.png", "Figure A - Platform Topology")
    add_image_if_exists(doc, repo / "documents/diagrams/orchestration-sequence.png", "Figure B - Orchestration Sequence")
    add_image_if_exists(
        doc,
        repo / "documents/diagrams/yaml-pluggability-control-surface.png",
        "Figure C - YAML Pluggability Control Surface",
    )

    heading(doc, "3) Component Stack and Service Endpoints", level=1)
    body(
        doc,
        "The following table merges the stack and service responsibilities from the system overview PDF with the current orchestration artifacts.",
    )
    add_component_table(doc)

    heading(doc, "4) Workflow State Machine and Decision Logic", level=1)
    body(
        doc,
        "The onboarding workflow transitions through deterministic stages (initialized -> KYC -> address -> AML -> credit -> risk -> completion), "
        "with retry envelopes and state history tracked per traceId.",
    )
    bullet(doc, "Escalate when confidence is below threshold, data is missing, or signals are contradictory.")
    bullet(doc, "Deny when policy conflict or high-risk conditions are detected.")
    bullet(doc, "Approve only when all critical controls and confidence conditions are satisfied.")

    heading(doc, "5) YAML Pluggability and Configuration Model", level=1)
    body(
        doc,
        "The platform is operationally pluggable through slot-based YAML configuration, supporting active version pointers, multi-version catalogs, "
        "transport-type abstraction (HTTP/local), and runtime safety controls.",
    )
    if yaml_deep_dive:
        for line in yaml_deep_dive:
            bullet(doc, line.lstrip("- ").strip())

    heading(doc, "6) Feature Highlights", level=1)
    if features:
        for line in features:
            bullet(doc, line)
    else:
        bullet(doc, "Event-native orchestration with deterministic state transitions.")
        bullet(doc, "Centralized decision governance and explainable outputs.")
        bullet(doc, "Trace-centric observability for diagnostics and compliance evidence.")

    heading(doc, "7) Implementation Anchors", level=1)
    if anchors:
        for line in anchors:
            bullet(doc, line.lstrip("- ").strip())

    heading(doc, "8) Source Integration Notes", level=1)
    body(doc, f"PDF pages parsed: {len(pdf_pages)}")
    body(doc, "The merged narrative incorporates: system purpose, architecture model, component stack, state-machine logic, and decision policy rules.")

    out_name = f"AGENTIC_ONBOARDING_MERGED_BRIEF_{datetime.now():%Y%m%d_%H%M%S}.docx"
    out_path = repo / out_name
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
